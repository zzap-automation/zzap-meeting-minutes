import asyncio
import os
import re
import tempfile
import uuid
from datetime import datetime
from io import BytesIO

import httpx
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse

# ── optional Supabase ────────────────────────────────────────────────────────
try:
    from supabase import create_client
    _sb_url = os.environ.get("SUPABASE_URL", "")
    _sb_key = os.environ.get("SUPABASE_KEY", "")
    supabase = create_client(_sb_url, _sb_key) if _sb_url and _sb_key else None
except Exception:
    supabase = None

# ── optional python-docx ─────────────────────────────────────────────────────
try:
    from docx import Document as DocxDoc
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ────────────────────────────────────────────────────────────────────────────
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_BASE    = "https://api.groq.com/openai/v1"
LLM_MODEL    = "llama-3.3-70b-versatile"
STT_MODEL    = "whisper-large-v3-turbo"

SYSTEM_PROMPT = """You are a professional meeting minutes writer for an architecture and planning firm.
Convert meeting transcripts into structured minutes using EXACTLY the format below.

# Meeting Minutes - [Concise Descriptive Title]

**Date:** [Extract from transcript, or "Not specified"]
**Attendees:** [Full names comma-separated, identified from transcript]
**Duration:** [Approximate, inferred from context, or "Not specified"]

---

## Meeting Discussion

| Agenda | Item | Description |
|--------|------|-------------|
| [Topic area — 1-4 words, repeated across rows sharing same broad topic] | [Specific sub-topic — 3-7 words] | [2-4 sentence factual prose. Third person. No bullet points inside cells.] |

---

## Action Items Summary

| Action Item | Details | Owner | Timeline |
|-------------|---------|-------|----------|
| [Short action title] | [What needs to be done] | [Person responsible] | [e.g. "This week", "ASAP", "TBD"] |

---

## Key Decisions

1. [Definitively agreed items only — not items to be discussed later]

---

## Additional Notes

- [Important context, background, or observations not captured in the tables]

RULES:
- Agenda: 1-4 words. Repeat across rows sharing the same broad topic.
- Item: 3-7 words, specific and scannable.
- Description: factual prose, complete sentences, third person voice.
- Action Items: every commitment or follow-up. Be specific on who, what, when.
- Key Decisions: definitively agreed items only.
- Additional Notes: context helping readers understand without having attended.
- Produce ALL five sections. Use "None identified." if a section has no content.
- Speaker labels like "Ife:" or "Stephen:" identify attendees — use them for ownership and names."""


SYSTEM_PROMPT_DETAILED = """You are a professional meeting minutes writer for an architecture and planning firm.
Convert meeting transcripts into detailed, comprehensive minutes using the format below.
Capture nuance, context, and reasoning — not just outcomes.

# Meeting Minutes - [Concise Descriptive Title]

**Date:** [Extract from transcript, or "Not specified"]
**Attendees:** [Full names comma-separated, identified from transcript]
**Duration:** [Approximate, inferred from context, or "Not specified"]

---

## Meeting Discussion

| Agenda | Item | Description |
|--------|------|-------------|
| [Topic area — 1-4 words, repeated across rows sharing same broad topic] | [Specific sub-topic — 3-7 words] | [4-6 sentences of detailed prose. Include context, rationale, differing views, and specifics discussed. Third person. No bullet points inside cells.] |

---

## Action Items Summary

| Action Item | Details | Owner | Timeline |
|-------------|---------|-------|----------|
| [Short action title] | [Specific details including any dependencies, constraints, or clarifications discussed] | [Person responsible] | [e.g. "This week", "ASAP", "TBD"] |

---

## Key Decisions

1. [Definitively agreed items — include brief rationale where discussed]

---

## Additional Notes

- [Important context, background, tangents, or observations that add understanding]

---

## Extended Summary

[Write 3-5 paragraphs of flowing prose summarising the full meeting in narrative form. Cover the key topics in the order they arose, the reasoning behind decisions, any disagreements or open questions, and the overall direction set. This section is for readers who want full context rather than just the structured highlights.]

RULES:
- Agenda: 1-4 words. Repeat across rows sharing the same broad topic.
- Item: 3-7 words, specific and scannable.
- Description: detailed prose, 4-6 complete sentences, third person voice.
- Include all six sections. Use "None identified." only if genuinely absent.
- Speaker labels like "Ife:" or "Stephen:" identify attendees.
- Extended Summary should read as a coherent narrative, not a list."""

# ── app ──────────────────────────────────────────────────────────────────────
app = FastAPI(title="zzap Meeting Minutes API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── helpers ──────────────────────────────────────────────────────────────────
def _groq_headers():
    if not GROQ_API_KEY:
        raise HTTPException(500, "GROQ_API_KEY is not configured.")
    return {"Authorization": f"Bearer {GROQ_API_KEY}"}


def _clean_md(text: str) -> str:
    """Strip accidental code fences from LLM output."""
    return re.sub(r"^```(?:markdown)?\n?", "", text.strip()).rstrip("`").strip()



def _truncate_transcript(text: str, max_chars: int = 28000):
    """Truncate a transcript to stay within Groq free-tier token limits.
    Keeps the first 60% and last 40% so opening context and closing
    decisions/action items are both preserved.
    Returns (text, was_truncated).
    """
    if len(text) <= max_chars:
        return text, False
    head = int(max_chars * 0.60)
    tail = max_chars - head
    truncated = (
        text[:head]
        + "\n\n[... transcript truncated to fit processing limits ...]\n\n"
        + text[-tail:]
    )
    return truncated, True


# ── routes ───────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "docx_available": DOCX_AVAILABLE, "supabase": supabase is not None}


@app.post("/transcribe")
async def transcribe(audio: UploadFile = File(...)):
    """Transcribe an uploaded audio file using Groq Whisper.

    If the file exceeds 24 MB, it is first compressed to 32 kbps mono MP3
    using ffmpeg (sufficient quality for speech, ~14 MB per hour of audio).
    """
    data = await audio.read()
    MAX_BYTES = 24 * 1024 * 1024  # stay safely under Groq's 25 MB limit

    if len(data) > MAX_BYTES:
        # Try ffmpeg compression to 32 kbps mono MP3
        try:
            with tempfile.NamedTemporaryFile(
                suffix=os.path.splitext(audio.filename or ".mp3")[1] or ".mp3",
                delete=False
            ) as src_f:
                src_f.write(data)
                src_path = src_f.name

            compressed_path = src_path + "_compressed.mp3"
            proc = await asyncio.create_subprocess_exec(
                "ffmpeg", "-y", "-i", src_path,
                "-ar", "16000",   # 16 kHz sample rate (Whisper native)
                "-ac", "1",       # mono
                "-b:a", "32k",    # 32 kbps — good for speech, ~14 MB/hr
                compressed_path,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.PIPE,
            )
            _, stderr = await proc.communicate()
            os.unlink(src_path)

            if proc.returncode != 0:
                raise RuntimeError(f"ffmpeg failed: {stderr.decode()}")

            with open(compressed_path, "rb") as f:
                data = f.read()
            os.unlink(compressed_path)

            if len(data) > MAX_BYTES:
                raise HTTPException(
                    413,
                    f"File is still {len(data)//1024//1024} MB after compression. "
                    "Please trim the recording to under ~90 minutes."
                )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Compression failed: {e}")

    ext = ".mp3"  # always MP3 after potential compression
    async with httpx.AsyncClient(timeout=300) as client:
        resp = await client.post(
            f"{GROQ_BASE}/audio/transcriptions",
            headers=_groq_headers(),
            files={"file": (f"audio{ext}", data, "audio/mpeg")},
            data={"model": STT_MODEL, "response_format": "text"},
        )
    if not resp.is_success:
        raise HTTPException(resp.status_code, f"Groq STT error: {resp.text}")
    return {"transcript": resp.text.strip()}


@app.post("/generate")
async def generate(
    transcript: str  = Form(...),
    title: str       = Form(""),
    attendees: str   = Form(""),
    date: str        = Form(""),
    detail_level: str = Form("standard"),   # "standard" | "detailed"
):
    """Generate formatted minutes from a transcript using Groq LLM."""
    transcript, was_truncated = _truncate_transcript(transcript)

    user_msg = "Please produce meeting minutes from this transcript.\n"
    if title:         user_msg += f"\nMeeting title: {title}"
    if attendees:     user_msg += f"\nAttendees: {attendees}"
    if date:          user_msg += f"\nDate: {date}"
    if was_truncated: user_msg += "\n\nNote: transcript was truncated due to length. Cover all sections present."
    user_msg += f"\n\nTRANSCRIPT:\n{transcript}"

    prompt = SYSTEM_PROMPT_DETAILED if detail_level == "detailed" else SYSTEM_PROMPT

    async with httpx.AsyncClient(timeout=180) as client:
        resp = await client.post(
            f"{GROQ_BASE}/chat/completions",
            headers={**_groq_headers(), "Content-Type": "application/json"},
            json={
                "model": LLM_MODEL,
                "messages": [
                    {"role": "system", "content": prompt},
                    {"role": "user",   "content": user_msg},
                ],
                "max_tokens": 6000,
                "temperature": 0.2,
            },
        )
    if not resp.is_success:
        raise HTTPException(resp.status_code, f"Groq LLM error: {resp.text}")
    minutes_md = _clean_md(resp.json()["choices"][0]["message"]["content"])
    return {"minutes_md": minutes_md}


@app.post("/save")
async def save_minutes(
    title: str      = Form(...),
    date: str       = Form(""),
    attendees: str  = Form(""),
    transcript: str = Form(""),
    minutes_md: str = Form(...),
    series: str     = Form(""),
):
    """Persist minutes to Supabase."""
    if supabase is None:
        raise HTTPException(503, "Supabase is not configured on this deployment.")
    result = supabase.table("meetings").insert({
        "title":      title,
        "date":       date,
        "attendees":  attendees,
        "transcript": transcript,
        "minutes_md": minutes_md,
        "series":     series.strip(),
    }).execute()
    return {"id": result.data[0]["id"]}


@app.get("/minutes")
async def list_minutes():
    """Return a summary list of all archived meetings."""
    if supabase is None:
        return {"meetings": []}
    result = (
        supabase.table("meetings")
        .select("id,title,date,attendees,series,created_at")
        .order("created_at", desc=True)
        .execute()
    )
    return {"meetings": result.data}


@app.get("/minutes/{meeting_id}")
async def get_minutes(meeting_id: str):
    """Return full record for one meeting."""
    if supabase is None:
        raise HTTPException(503, "Supabase is not configured.")
    result = supabase.table("meetings").select("*").eq("id", meeting_id).execute()
    if not result.data:
        raise HTTPException(404, "Meeting not found.")
    return result.data[0]


@app.get("/series")
async def list_series():
    """Return all unique non-empty series names."""
    if supabase is None:
        return {"series": []}
    result = supabase.table("meetings").select("series").execute()
    names = sorted({r["series"].strip() for r in result.data if r.get("series","").strip()})
    return {"series": names}


@app.post("/query")
async def query_minutes(
    question: str = Form(...),
    series: str   = Form(""),
):
    """Answer a question by searching the meeting archive and synthesising with Groq."""
    if supabase is None:
        raise HTTPException(503, "Supabase is not configured.")

    relevant = []
    try:
        q = supabase.table("meetings").select(
            "id,title,date,attendees,minutes_md,series"
        )
        if series.strip():
            q = q.eq("series", series.strip())
        fts_result = q.text_search("search_vector", question).limit(6).execute()
        relevant = fts_result.data or []
    except Exception:
        pass

    if not relevant:
        try:
            q2 = supabase.table("meetings").select(
                "id,title,date,attendees,minutes_md,series"
            ).order("created_at", desc=True)
            if series.strip():
                q2 = q2.eq("series", series.strip())
            fallback = q2.limit(8).execute()
            relevant = fallback.data or []
        except Exception:
            pass

    if not relevant:
        return {
            "answer": "No meeting minutes found in the archive to search through.",
            "sources": [],
        }

    context_parts = []
    for m in relevant:
        title     = m.get("title", "Untitled")
        date      = m.get("date", "unknown")
        attendees = m.get("attendees", "")
        body      = m.get("minutes_md", "")
        header    = "Meeting: " + title + " | Date: " + date + " | Attendees: " + attendees
        context_parts.append("--- " + header + " ---\n" + body)
    context = "\n\n".join(context_parts)

    system = (
        "You are an assistant that answers questions about meeting minutes for an "
        "architecture and planning firm. Answer ONLY from the provided meeting minutes. "
        "Be specific and concise. For each key piece of information, cite the meeting "
        "title and date in parentheses, e.g. (Weekly Check-in, April 7 2026). "
        "If the answer is not found in the provided meetings, say so clearly."
    )
    user_msg = "Meeting minutes archive:\n\n" + context + "\n\nQuestion: " + question

    async with httpx.AsyncClient(timeout=60) as client:
        resp = await client.post(
            f"{GROQ_BASE}/chat/completions",
            headers={**_groq_headers(), "Content-Type": "application/json"},
            json={
                "model": LLM_MODEL,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user",   "content": user_msg},
                ],
                "max_tokens": 1000,
                "temperature": 0.1,
            },
        )
    if not resp.is_success:
        raise HTTPException(resp.status_code, f"Groq error: {resp.text}")

    answer  = resp.json()["choices"][0]["message"]["content"].strip()
    sources = [
        {"id": m["id"], "title": m.get("title", "Untitled"), "date": m.get("date", "")}
        for m in relevant
    ]
    return {"answer": answer, "sources": sources}


@app.delete("/minutes/{meeting_id}")
async def delete_minutes(meeting_id: str):
    """Delete a meeting from the archive."""
    if supabase is None:
        raise HTTPException(503, "Supabase is not configured.")
    result = supabase.table("meetings").delete().eq("id", meeting_id).execute()
    if not result.data:
        raise HTTPException(404, "Meeting not found.")
    return {"deleted": meeting_id}


@app.patch("/minutes/{meeting_id}")
async def update_minutes(
    meeting_id: str,
    title:     str = Form(""),
    date:      str = Form(""),
    attendees: str = Form(""),
    series:    str = Form(""),
):
    """Update editable fields on a saved meeting."""
    if supabase is None:
        raise HTTPException(503, "Supabase is not configured.")
    result = (
        supabase.table("meetings")
        .update({
            "title":     title.strip(),
            "date":      date.strip(),
            "attendees": attendees.strip(),
            "series":    series.strip(),
        })
        .eq("id", meeting_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(404, "Meeting not found.")
    return result.data[0]


@app.post("/series/rename")
async def rename_series(
    old_name: str = Form(...),
    new_name: str = Form(...),
):
    """Rename a series across all meetings that use it."""
    if supabase is None:
        raise HTTPException(503, "Supabase is not configured.")
    old_name = old_name.strip()
    new_name = new_name.strip()
    if not old_name:
        raise HTTPException(400, "old_name is required.")
    result = (
        supabase.table("meetings")
        .update({"series": new_name})
        .eq("series", old_name)
        .execute()
    )
    updated = len(result.data) if result.data else 0
    return {"updated": updated, "old_name": old_name, "new_name": new_name}


@app.post("/export/docx")
async def export_docx(minutes_md: str = Form(...), filename: str = Form("meeting-minutes")):
    """Convert markdown minutes to a formatted .docx file."""
    if not DOCX_AVAILABLE:
        raise HTTPException(503, "python-docx is not installed on this server.")

    buf = _md_to_docx(minutes_md)
    safe_name = re.sub(r"[^\w\-]", "-", filename).strip("-") or "meeting-minutes"
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )


# ── markdown → docx converter ────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str):
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _add_bottom_border(paragraph, color_hex: str = "CCCCCC"):
    """Draw a thin bottom rule under a paragraph (used after section headers)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _parse_inline(text: str):
    """Return list of (text, bold) tuples from a markdown inline string."""
    parts, buf, bold = [], "", False
    i = 0
    while i < len(text):
        if text[i:i+2] == "**":
            parts.append((buf, bold)); buf = ""; bold = not bold; i += 2
        else:
            buf += text[i]; i += 1
    parts.append((buf, bold))
    return [(t, b) for t, b in parts if t]


def _add_paragraph_runs(para, text: str, base_bold=False, base_size_pt=11):
    for chunk, chunk_bold in _parse_inline(text):
        run = para.add_run(chunk)
        run.bold = base_bold or chunk_bold
        run.font.size = Pt(base_size_pt)


def _md_to_docx(md: str) -> BytesIO:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not available")

    doc = DocxDoc()

    # Word 2013+ compat (prevents "Compatibility Mode" banner on open)
    settings_el = doc.settings.element
    compat_el   = OxmlElement("w:compat")
    cs          = OxmlElement("w:compatSetting")
    cs.set(qn("w:name"), "compatibilityMode")
    cs.set(qn("w:uri"),  "http://schemas.microsoft.com/office/word")
    cs.set(qn("w:val"),  "15")
    compat_el.append(cs)
    settings_el.append(compat_el)

    # Page size US Letter + 1-inch margins
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Set Word to open in Print Layout, zoomed to show one full page
    try:
        zoom_el = OxmlElement("w:zoom")
        zoom_el.set(qn("w:val"), "fullPage")   # = "One Page" in Word's View menu
        doc.settings.element.insert(0, zoom_el)
        view_el = OxmlElement("w:view")
        view_el.set(qn("w:val"), "print")
        doc.settings.element.insert(0, view_el)
    except Exception:
        pass

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines     = md.split("\n")
    i         = 0
    in_table  = False

    # We collect table rows then flush them to a real docx table
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        col_count = max(len(r) for r in table_rows)
        # Column widths (DXA): content width = 9360
        if col_count == 3:
            widths = [1800, 2200, 5360]
        elif col_count == 4:
            widths = [1800, 3200, 1800, 2560]
        else:
            each = 9360 // col_count
            widths = [each] * col_count

        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _el

        # ── pre-identify agenda merge groups ────────────────────────────────
        # Build a lookup: row_idx -> ('restart' | 'continue' | None)
        merge_role = {}   # row_idx: 'restart' or 'continue'
        if len(table_rows) > 2:
            grp_start = 1
            grp_text  = table_rows[1][0].strip() if len(table_rows) > 1 else ""
            for ri in range(2, len(table_rows)):
                ct = table_rows[ri][0].strip() if table_rows[ri] else ""
                if ct != grp_text:
                    if ri - grp_start > 1:
                        merge_role[grp_start] = 'restart'
                        for ci in range(grp_start + 1, ri):
                            merge_role[ci] = 'continue'
                    grp_start, grp_text = ri, ct
            last = len(table_rows) - 1
            if last - grp_start >= 1:
                merge_role[grp_start] = 'restart'
                for ci in range(grp_start + 1, last + 1):
                    merge_role[ci] = 'continue'

        tbl = doc.add_table(rows=0, cols=col_count)
        tbl.style = "Table Grid"

        for row_idx, cells in enumerate(table_rows):
            row = tbl.add_row()
            for col_idx in range(col_count):
                cell = row.cells[col_idx]
                # For continuation rows, col 0 should be empty
                if col_idx == 0 and merge_role.get(row_idx) == 'continue':
                    text = ""
                else:
                    text = cells[col_idx] if col_idx < len(cells) else ""
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # ── vMerge tag (col 0 only, inline during construction) ────
                if col_idx == 0 and row_idx in merge_role:
                    vm = _el('w:vMerge')
                    if merge_role[row_idx] == 'restart':
                        vm.set(_qn('w:val'), 'restart')
                    tcPr.append(vm)

                # ── cell width ────────────────────────────────────────────
                tcW = _el("w:tcW")
                tcW.set(_qn("w:w"),    str(widths[col_idx]))
                tcW.set(_qn("w:type"), "dxa")
                tcPr.append(tcW)

                # ── content ───────────────────────────────────────────────
                if row_idx == 0:                          # header
                    _set_cell_shading(cell, "1565a8")
                    run = cell.paragraphs[0].add_run(text)
                    run.bold           = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size      = Pt(10)
                elif merge_role.get(row_idx) == 'continue' and col_idx == 0:
                    pass  # leave continuation cell empty
                else:
                    _add_paragraph_runs(cell.paragraphs[0], text, base_size_pt=10)
                    if row_idx % 2 == 0:
                        _set_cell_shading(cell, "F5F5F5")

        doc.add_paragraph()
        in_table   = False
        table_rows.clear()

    while i < len(lines):
        line = lines[i]

        # ── H1 title ──────────────────────────────────────────────────────
        if line.startswith("# "):
            flush_table()
            text = line[2:].strip()
            p    = doc.add_paragraph()
            run  = p.add_run(text)
            run.bold           = True
            run.font.size      = Pt(18)
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xa8)  # dark navy
            p.paragraph_format.space_after = Pt(4)
            i += 1; continue

        # ── H2 section header ─────────────────────────────────────────────
        if line.startswith("## "):
            flush_table()
            text = line[3:].strip()
            p    = doc.add_paragraph()
            run  = p.add_run(text.upper())
            run.bold           = True
            run.font.size      = Pt(10)
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xa8)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(2)
            _add_bottom_border(p, "1565a8")
            i += 1; continue

        # ── horizontal rule ───────────────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            flush_table()
            i += 1; continue

        # ── meta lines (**Date:** etc.) ───────────────────────────────────
        if line.startswith("**") and ":**" in line:
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            _add_paragraph_runs(p, line)
            i += 1; continue

        # ── table row ─────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            in_table = True
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # skip separator rows
            if all(re.match(r"^[-: ]+$", c) for c in cells if c):
                i += 1; continue
            table_rows.append(cells)
            i += 1; continue
        elif in_table:
            flush_table()

        # ── numbered list ─────────────────────────────────────────────────
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            flush_table()
            p    = doc.add_paragraph(style="List Number")
            _add_paragraph_runs(p, m.group(1))
            p.paragraph_format.space_after = Pt(2)
            i += 1; continue

        # ── bullet list ───────────────────────────────────────────────────
        if re.match(r"^[-*]\s+", line):
            flush_table()
            text = re.sub(r"^[-*]\s+", "", line)
            p    = doc.add_paragraph(style="List Bullet")
            _add_paragraph_runs(p, text)
            p.paragraph_format.space_after = Pt(2)
            i += 1; continue

        # ── blank line ────────────────────────────────────────────────────
        if not line.strip():
            flush_table()
            i += 1; continue

        # ── plain paragraph ───────────────────────────────────────────────
        flush_table()
        p = doc.add_paragraph()
        _add_paragraph_runs(p, line)
        i += 1

    flush_table()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
